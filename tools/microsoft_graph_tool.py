"""Microsoft Graph helper functions for Hermes."""

from __future__ import annotations

import json
import urllib.parse
import urllib.request
from typing import Any

from agent.microsoft_oauth import get_access_token


GRAPH_BASE_URL = "https://graph.microsoft.com/v1.0"


def _request(method: str, path: str, payload: dict[str, Any] | None = None, params: dict[str, Any] | None = None) -> dict[str, Any]:
    token = get_access_token()

    url = GRAPH_BASE_URL + path
    if params:
        url += "?" + urllib.parse.urlencode(params)

    data = None
    headers = {
        "Authorization": f"Bearer {token}",
        "Accept": "application/json",
    }

    if payload is not None:
        data = json.dumps(payload).encode("utf-8")
        headers["Content-Type"] = "application/json"

    req = urllib.request.Request(url, data=data, headers=headers, method=method)

    with urllib.request.urlopen(req, timeout=30) as resp:
        raw = resp.read().decode("utf-8")
        return json.loads(raw) if raw else {}


def graph_get(path: str, params: dict[str, Any] | None = None) -> dict[str, Any]:
    return _request("GET", path, params=params)


def graph_post(path: str, payload: dict[str, Any]) -> dict[str, Any]:
    return _request("POST", path, payload=payload)


def graph_patch(path: str, payload: dict[str, Any]) -> dict[str, Any]:
    return _request("PATCH", path, payload=payload)


def graph_delete(path: str) -> dict[str, Any]:
    return _request("DELETE", path)


def list_recent_messages(limit: int = 10) -> list[dict[str, Any]]:
    data = graph_get(
        "/me/messages",
        {
            "$top": max(1, min(int(limit), 50)),
            "$select": "id,subject,from,receivedDateTime,webLink",
            "$orderby": "receivedDateTime DESC",
        },
    )
    return data.get("value", [])


def get_message(message_id: str) -> dict[str, Any]:
    if not message_id:
        raise ValueError("message_id is required")

    return graph_get(
        f"/me/messages/{urllib.parse.quote(message_id)}",
        {
            "$select": "id,subject,from,toRecipients,ccRecipients,receivedDateTime,body,bodyPreview,webLink",
        },
    )


def search_messages(query: str, limit: int = 10) -> list[dict[str, Any]]:
    if not query or not query.strip():
        raise ValueError("query is required")

    data = graph_get(
        "/me/messages",
        {
            "$top": max(1, min(int(limit), 50)),
            "$search": f'"{query.strip()}"',
            "$select": "id,subject,from,receivedDateTime,bodyPreview,webLink",
        },
    )
    return data.get("value", [])


def _recipient_list(to: str | list[str]) -> list[dict[str, dict[str, str]]]:
    if isinstance(to, str):
        recipients = [addr.strip() for addr in to.split(",") if addr.strip()]
    else:
        recipients = [str(addr).strip() for addr in to if str(addr).strip()]

    if not recipients:
        raise ValueError("At least one recipient is required")

    return [{"emailAddress": {"address": address}} for address in recipients]


def send_message(to: str | list[str], subject: str, body: str, save_to_sent_items: bool = True) -> dict[str, Any]:
    if not subject:
        raise ValueError("subject is required")
    if not body:
        raise ValueError("body is required")

    payload = {
        "message": {
            "subject": subject,
            "body": {"contentType": "Text", "content": body},
            "toRecipients": _recipient_list(to),
        },
        "saveToSentItems": save_to_sent_items,
    }

    return graph_post("/me/sendMail", payload)


def create_draft(to: str | list[str], subject: str, body: str) -> dict[str, Any]:
    if not subject:
        raise ValueError("subject is required")
    if not body:
        raise ValueError("body is required")

    payload = {
        "subject": subject,
        "body": {"contentType": "Text", "content": body},
        "toRecipients": _recipient_list(to),
    }

    return graph_post("/me/messages", payload)


def summarizable_message_text(message_id: str) -> str:
    msg = get_message(message_id)

    sender = (msg.get("from") or {}).get("emailAddress", {})
    body = msg.get("body") or {}

    parts = [
        f"Subject: {msg.get('subject') or ''}",
        f"From: {sender.get('name') or ''} <{sender.get('address') or ''}>",
        f"Received: {msg.get('receivedDateTime') or ''}",
        "",
        body.get("content") or msg.get("bodyPreview") or "",
    ]

    return "\n".join(parts).strip()


def _html_to_text(html: str) -> str:
    import re
    from html import unescape

    text = re.sub(r"(?is)<(script|style).*?>.*?</\1>", " ", html or "")
    text = re.sub(r"(?i)<br\s*/?>", "\n", text)
    text = re.sub(r"(?i)</p\s*>", "\n", text)
    text = re.sub(r"(?i)</div\s*>", "\n", text)
    text = re.sub(r"<[^>]+>", " ", text)
    text = unescape(text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n\s*\n+", "\n\n", text)
    return text.strip()


def summarizable_message_text(message_id: str) -> str:
    msg = get_message(message_id)

    sender = (msg.get("from") or {}).get("emailAddress", {})
    body = msg.get("body") or {}
    raw_body = body.get("content") or msg.get("bodyPreview") or ""

    content_type = (body.get("contentType") or "").lower()
    clean_body = _html_to_text(raw_body) if content_type == "html" else raw_body

    parts = [
        f"Subject: {msg.get('subject') or ''}",
        f"From: {sender.get('name') or ''} <{sender.get('address') or ''}>",
        f"Received: {msg.get('receivedDateTime') or ''}",
        "",
        clean_body,
    ]

    return "\n".join(parts).strip()


def list_calendar_events(limit: int = 10) -> list[dict[str, Any]]:
    data = graph_get(
        "/me/events",
        {
            "$top": max(1, min(int(limit), 50)),
            "$select": "id,subject,start,end,location,organizer,attendees,webLink",
            "$orderby": "start/dateTime",
        },
    )
    return data.get("value", [])


def list_upcoming_calendar_events(limit: int = 10, days_ahead: int = 30) -> list[dict[str, Any]]:
    from datetime import datetime, timedelta, timezone

    now = datetime.now(timezone.utc)
    end = now + timedelta(days=max(1, int(days_ahead)))

    data = graph_get(
        "/me/calendarView",
        {
            "startDateTime": now.isoformat(),
            "endDateTime": end.isoformat(),
            "$top": max(1, min(int(limit), 50)),
            "$select": "id,subject,start,end,location,organizer,attendees,webLink",
            "$orderby": "start/dateTime",
        },
    )
    return data.get("value", [])


def create_calendar_event(
    subject: str,
    start_datetime: str,
    end_datetime: str,
    timezone: str = "America/New_York",
    location: str = "",
    body: str = "",
    attendees: list[str] | None = None,
) -> dict[str, Any]:
    if not subject:
        raise ValueError("subject is required")
    if not start_datetime:
        raise ValueError("start_datetime is required")
    if not end_datetime:
        raise ValueError("end_datetime is required")

    payload = {
        "subject": subject,
        "body": {
            "contentType": "Text",
            "content": body or "",
        },
        "start": {
            "dateTime": start_datetime,
            "timeZone": timezone,
        },
        "end": {
            "dateTime": end_datetime,
            "timeZone": timezone,
        },
        "location": {
            "displayName": location or "",
        },
        "attendees": [
            {
                "emailAddress": {
                    "address": email,
                },
                "type": "required",
            }
            for email in (attendees or [])
        ],
    }

    return graph_post("/me/events", payload)


def get_calendar_event(event_id: str) -> dict[str, Any]:
    if not event_id:
        raise ValueError("event_id is required")

    return graph_get(
        f"/me/events/{urllib.parse.quote(event_id)}",
        {
            "$select": "id,subject,start,end,location,organizer,attendees,body,bodyPreview,webLink",
        },
    )


def update_calendar_event(
    event_id: str,
    subject: str | None = None,
    start_datetime: str | None = None,
    end_datetime: str | None = None,
    timezone: str = "America/New_York",
    location: str | None = None,
    body: str | None = None,
    attendees: list[str] | None = None,
) -> dict[str, Any]:
    if not event_id:
        raise ValueError("event_id is required")

    payload: dict[str, Any] = {}

    if subject is not None:
        payload["subject"] = subject

    if body is not None:
        payload["body"] = {
            "contentType": "Text",
            "content": body,
        }

    if location is not None:
        payload["location"] = {
            "displayName": location,
        }

    if start_datetime is not None:
        payload["start"] = {
            "dateTime": start_datetime,
            "timeZone": timezone,
        }

    if end_datetime is not None:
        payload["end"] = {
            "dateTime": end_datetime,
            "timeZone": timezone,
        }

    if attendees is not None:
        payload["attendees"] = [
            {
                "emailAddress": {
                    "address": email,
                },
                "type": "required",
            }
            for email in attendees
        ]

    if not payload:
        raise ValueError("No update fields provided")

    return graph_patch(f"/me/events/{urllib.parse.quote(event_id)}", payload)


def delete_calendar_event(event_id: str) -> bool:
    if not event_id:
        raise ValueError("event_id is required")

    graph_delete(f"/me/events/{urllib.parse.quote(event_id)}")
    return True



def list_onedrive_root(limit: int = 25) -> list[dict[str, Any]]:
    data = graph_get(
        "/me/drive/root/children",
        {
            "$top": max(1, min(int(limit), 100)),
            "$select": "id,name,size,webUrl,file,folder,lastModifiedDateTime",
        },
    )
    return data.get("value", [])


def graph_put_bytes(path: str, content: bytes, content_type: str = "application/octet-stream") -> dict[str, Any]:
    token = get_access_token()

    url = GRAPH_BASE_URL + path

    req = urllib.request.Request(
        url,
        data=content,
        headers={
            "Authorization": f"Bearer {token}",
            "Accept": "application/json",
            "Content-Type": content_type,
        },
        method="PUT",
    )

    with urllib.request.urlopen(req, timeout=30) as resp:
        raw = resp.read().decode("utf-8")
        return json.loads(raw) if raw else {}


def upload_text_file_to_onedrive(filename: str, content: str) -> dict[str, Any]:
    if not filename:
        raise ValueError("filename is required")

    safe_name = filename.strip().replace("/", "-")
    upload_path = f"/me/drive/root:/{urllib.parse.quote(safe_name)}:/content"

    return graph_put_bytes(
        upload_path,
        content.encode("utf-8"),
        "text/plain; charset=utf-8",
    )


def list_todo_lists() -> list[dict[str, Any]]:
    data = graph_get(
        "/me/todo/lists",
        {
            "$select": "id,displayName,isOwner,isShared,isDefaultList",
        },
    )
    return data.get("value", [])


def create_word_docx_and_upload(filename: str, title: str, body: str) -> dict[str, Any]:
    from io import BytesIO
    from docx import Document

    if not filename:
        raise ValueError("filename is required")
    if not filename.lower().endswith(".docx"):
        filename = filename + ".docx"

    doc = Document()

    if title:
        doc.add_heading(title, level=1)

    for paragraph in (body or "").split("\n\n"):
        clean = paragraph.strip()
        if clean:
            doc.add_paragraph(clean)

    buffer = BytesIO()
    doc.save(buffer)

    safe_name = filename.strip().replace("/", "-")
    upload_path = f"/me/drive/root:/{urllib.parse.quote(safe_name)}:/content"

    return graph_put_bytes(
        upload_path,
        buffer.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def list_todo_lists() -> list[dict[str, Any]]:
    data = graph_get("/me/todo/lists")
    return data.get("value", [])


def create_todo_task(
    title: str,
    body: str = "",
    list_id: str | None = None,
    due_datetime: str | None = None,
    timezone: str = "America/New_York",
) -> dict[str, Any]:
    if not title:
        raise ValueError("title is required")

    if not list_id:
        lists = list_todo_lists()
        if not lists:
            raise RuntimeError("No Microsoft To Do lists found.")
        list_id = lists[0]["id"]

    payload: dict[str, Any] = {
        "title": title,
    }

    if body:
        payload["body"] = {
            "content": body,
            "contentType": "text",
        }

    if due_datetime:
        payload["dueDateTime"] = {
            "dateTime": due_datetime,
            "timeZone": timezone,
        }

    return graph_post(f"/me/todo/lists/{list_id}/tasks", payload)


def list_todo_tasks(list_id: str | None = None, limit: int = 25) -> list[dict[str, Any]]:
    if not list_id:
        lists = list_todo_lists()
        if not lists:
            raise RuntimeError("No Microsoft To Do lists found.")
        list_id = lists[0]["id"]

    data = graph_get(
        f"/me/todo/lists/{list_id}/tasks",
        {
            "$top": max(1, min(int(limit), 100)),
        },
    )
    return data.get("value", [])


def update_todo_task(
    task_id: str,
    list_id: str | None = None,
    title: str | None = None,
    body: str | None = None,
    status: str | None = None,
    due_datetime: str | None = None,
    timezone: str = "America/New_York",
) -> dict[str, Any]:
    if not task_id:
        raise ValueError("task_id is required")

    if not list_id:
        lists = list_todo_lists()
        if not lists:
            raise RuntimeError("No Microsoft To Do lists found.")
        list_id = lists[0]["id"]

    payload: dict[str, Any] = {}

    if title is not None:
        payload["title"] = title

    if body is not None:
        payload["body"] = {
            "content": body,
            "contentType": "text",
        }

    if status is not None:
        payload["status"] = status

    if due_datetime is not None:
        payload["dueDateTime"] = {
            "dateTime": due_datetime,
            "timeZone": timezone,
        }

    if not payload:
        raise ValueError("No update fields provided")

    return graph_patch(f"/me/todo/lists/{list_id}/tasks/{task_id}", payload)


def delete_todo_task(task_id: str, list_id: str | None = None) -> bool:
    if not task_id:
        raise ValueError("task_id is required")

    if not list_id:
        lists = list_todo_lists()
        if not lists:
            raise RuntimeError("No Microsoft To Do lists found.")
        list_id = lists[0]["id"]

    graph_delete(f"/me/todo/lists/{list_id}/tasks/{task_id}")
    return True


# === Normalized Microsoft To Do helpers ===

def _todo_lists_by_name() -> dict[str, dict[str, Any]]:
    return {
        str(item.get("displayName") or "").strip().lower(): item
        for item in list_todo_lists()
    }


def get_todo_list_id_by_name(name: str = "Tasks") -> str:
    if not name:
        name = "Tasks"

    lists = _todo_lists_by_name()
    wanted = name.strip().lower()

    if wanted in lists:
        return str(lists[wanted]["id"])

    if "tasks" in lists:
        return str(lists["tasks"]["id"])

    if lists:
        return str(next(iter(lists.values()))["id"])

    raise RuntimeError("No Microsoft To Do lists found.")


def list_todo_tasks_by_list_name(list_name: str = "Tasks", limit: int = 50) -> list[dict[str, Any]]:
    list_id = get_todo_list_id_by_name(list_name)
    return list_todo_tasks(list_id, limit=limit)


def find_todo_task_by_title(title: str, list_name: str = "Tasks", include_completed: bool = True) -> dict[str, Any] | None:
    if not title or not title.strip():
        raise ValueError("title is required")

    query = title.strip().lower()
    tasks = list_todo_tasks_by_list_name(list_name, limit=100)

    for task in tasks:
        task_title = str(task.get("title") or "").strip().lower()
        status = str(task.get("status") or "").strip().lower()

        if not include_completed and status == "completed":
            continue

        if task_title == query:
            task["todo_list_id"] = get_todo_list_id_by_name(list_name)
            return task

    for task in tasks:
        task_title = str(task.get("title") or "").strip().lower()
        status = str(task.get("status") or "").strip().lower()

        if not include_completed and status == "completed":
            continue

        if query in task_title:
            task["todo_list_id"] = get_todo_list_id_by_name(list_name)
            return task

    return None


def create_todo_task_by_list_name(
    title: str,
    body: str = "",
    list_name: str = "Tasks",
    due_datetime: str | None = None,
    timezone: str = "America/New_York",
) -> dict[str, Any]:
    if not title:
        raise ValueError("title is required")

    list_id = get_todo_list_id_by_name(list_name)

    payload: dict[str, Any] = {
        "title": title,
    }

    if body:
        payload["body"] = {
            "content": body,
            "contentType": "text",
        }

    if due_datetime:
        payload["dueDateTime"] = {
            "dateTime": due_datetime,
            "timeZone": timezone,
        }

    return graph_post(f"/me/todo/lists/{list_id}/tasks", payload)


def complete_todo_task_by_title(title: str, list_name: str = "Tasks") -> dict[str, Any]:
    task = find_todo_task_by_title(title, list_name=list_name, include_completed=True)
    if not task:
        raise RuntimeError(f"Microsoft To Do task not found: {title!r} in list {list_name!r}")

    list_id = str(task.get("todo_list_id") or get_todo_list_id_by_name(list_name))
    task_id = str(task["id"])

    return graph_patch(
        f"/me/todo/lists/{list_id}/tasks/{task_id}",
        {"status": "completed"},
    )


def delete_todo_task_by_title(title: str, list_name: str = "Tasks") -> bool:
    task = find_todo_task_by_title(title, list_name=list_name, include_completed=True)
    if not task:
        raise RuntimeError(f"Microsoft To Do task not found: {title!r} in list {list_name!r}")

    list_id = str(task.get("todo_list_id") or get_todo_list_id_by_name(list_name))
    task_id = str(task["id"])

    graph_delete(f"/me/todo/lists/{list_id}/tasks/{task_id}")
    return True

